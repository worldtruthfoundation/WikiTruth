import io
import streamlit as st
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

def create_wiki_document(article, language_name, is_translated=False, translated_to=None):
    """
    Creates a DOCX document from Wikipedia article content
    
    Args:
        article (dict): Wikipedia article data containing title, summary, content, and URL
        language_name (str): Name of the language the article is in
        is_translated (bool): Whether this is a translated version
        translated_to (str): Name of the language the article was translated to
        
    Returns:
        BytesIO: Document as a bytes buffer
    """
    doc = Document()
    
    # Set document properties
    core_properties = doc.core_properties
    core_properties.title = article["title"]
    core_properties.subject = f"Wikipedia Article - {'Original' if not is_translated else 'Translation'}"
    
    # Add title
    title = doc.add_heading(article["title"], level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add source information
    source_para = doc.add_paragraph()
    source_text = f"Source: {article['url']}"
    source_run = source_para.add_run(source_text)
    source_run.italic = True
    source_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add language information
    lang_para = doc.add_paragraph()
    if is_translated:
        lang_text = f"Original: {language_name} | Translated to: {translated_to}"
    else:
        lang_text = f"Language: {language_name}"
    lang_run = lang_para.add_run(lang_text)
    lang_run.italic = True
    lang_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add horizontal line
    doc.add_paragraph("_" * 50)
    
    # Add summary section
    doc.add_heading("Summary", level=1)
    doc.add_paragraph(article["summary"])
    
    # Add content
    doc.add_heading("Full Content", level=1)
    
    # Split content into sections based on headings and format appropriately
    sections = split_content_into_sections(article["content"])
    
    for section in sections:
        # Determine heading level based on the number of '=' signs
        if section["title"]:
            level = min(section["level"], 5)  # Cap at heading level 5
            doc.add_heading(section["title"], level=level)
        
        # Add the content
        paragraphs = section["content"].split("\n\n")
        for paragraph in paragraphs:
            if paragraph.strip():
                doc.add_paragraph(paragraph.strip())
    
    # Save to bytes buffer
    doc_buffer = io.BytesIO()
    doc.save(doc_buffer)
    doc_buffer.seek(0)
    
    return doc_buffer

def split_content_into_sections(content):
    """
    Split Wikipedia content into sections based on headings for better document structuring
    
    Args:
        content (str): Wikipedia article content
        
    Returns:
        list: List of dictionaries with section titles and content
    """
    # Find all section headings using regex
    # This pattern matches heading patterns like "== Title ==" or "=== Subsection ==="
    heading_pattern = re.compile(r'^(={2,6})\s*(.*?)\s*\1', re.MULTILINE)
    
    # Find all headings and their positions
    headings = []
    for match in heading_pattern.finditer(content):
        level = len(match.group(1))
        title = match.group(2)
        pos = match.start()
        headings.append((pos, level, title))
    
    # Sort headings by position
    headings.sort()
    
    # Create sections
    sections = []
    
    # Add the initial section (before the first heading)
    if headings:
        first_pos = headings[0][0]
        intro_content = content[:first_pos].strip()
        if intro_content:
            sections.append({
                "title": None,
                "content": intro_content,
                "level": 0
            })
    
    # Process each heading and its content
    for i in range(len(headings)):
        pos, level, title = headings[i]
        
        # Calculate where this section ends
        if i < len(headings) - 1:
            end_pos = headings[i+1][0]
        else:
            end_pos = len(content)
        
        # Get the heading match to skip it
        heading_match = heading_pattern.search(content[pos:end_pos])
        if heading_match:
            section_content = content[pos + heading_match.end():end_pos].strip()
        else:
            section_content = content[pos:end_pos].strip()
        
        sections.append({
            "title": title,
            "content": section_content,
            "level": level // 2  # Convert to Word heading level (1-5)
        })
    
    # If no headings were found, return the entire content as one section
    if not sections:
        sections.append({
            "title": None,
            "content": content,
            "level": 0
        })
    
    return sections

def get_download_filename(article_title, is_translated=False, language_code=None):
    """
    Generate a filename for the document download
    
    Args:
        article_title (str): Title of the article
        is_translated (bool): Whether this is a translated version
        language_code (str): Language code if applicable
        
    Returns:
        str: Properly formatted filename
    """
    # Clean title for use in filename - keep it more readable
    # Remove special characters that aren't allowed in filenames
    clean_title = article_title.replace("/", "-").replace("\\", "-")
    clean_title = re.sub(r'[<>:"|?*]', '', clean_title)
    # Replace multiple spaces with a single underscore for readability
    clean_title = re.sub(r'\s+', '_', clean_title)
    
    if is_translated:
        return f"Wikipedia-{clean_title}-translated_to_{language_code}.docx"
    else:
        return f"Wikipedia-{clean_title}.docx"
